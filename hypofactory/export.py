"""Экспорт результатов: бизнес-отчёт DOCX, автономный HTML, задачи CSV/JSON."""
from __future__ import annotations

import csv
import html as html_mod
import io
import json

from .pipeline import PipelineResult

CRITERIA_RU = {
    "value": "Ценность",
    "feasibility": "Реализуемость",
    "novelty": "Новизна",
    "evidence": "Обоснованность",
    "testability": "Проверяемость",
    "feedback": "Фидбэк экспертов",
}


# ------------------------------------------------------------------ DOCX
def export_docx(result: PipelineResult) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("Фабрика гипотез — отчёт по генерации исследовательской повестки", 0)

    p = doc.add_paragraph()
    p.add_run("Цель: ").bold = True
    p.add_run(result.goal)
    p = doc.add_paragraph()
    p.add_run("Ограничения: ").bold = True
    p.add_run(result.constraints)
    p = doc.add_paragraph()
    p.add_run("Источник данных: ").bold = True
    p.add_run(result.report.source_file)

    doc.add_heading("1. Диагностика потерь (факты из отчёта института)", level=1)
    for s in result.signals:
        doc.add_heading(f"{s.title} — {s.tonnes:,.0f} т {s.metal}", level=2)
        doc.add_paragraph(s.explanation)

    doc.add_heading("2. Ранжированные гипотезы", level=1)
    for h in result.hypotheses:
        doc.add_heading(f"№{h.rank}. {h.statement}", level=2)
        t = doc.add_paragraph()
        t.add_run(f"Итоговый балл: {h.total_score:.2f}   |   "
                  f"Сигнал: {h.signal_title} ({h.signal_tonnes:,.0f} т {h.metal})").italic = True

        doc.add_paragraph(f"Механизм: {h.mechanism}")
        doc.add_paragraph(f"Обоснование: {h.rationale}")
        eff = h.expected_effect
        doc.add_paragraph(
            f"Ожидаемый эффект: возврат {eff['tonnes_min']:,.0f}–"
            f"{eff['tonnes_max']:,.0f} т {h.metal} за период "
            f"({eff['min_recovery_pct']:.0f}–{eff['max_recovery_pct']:.0f}% "
            f"{eff['basis']})")

        if h.scores:
            doc.add_paragraph("Разбор оценки:")
            for crit, sc in h.scores.items():
                doc.add_paragraph(
                    f"{crit}: {sc['score']:.2f} (вес {sc.get('weight') or 0:.2f}) — "
                    f"{sc['explanation']}", style="List Bullet")
        if h.risks:
            doc.add_paragraph("Риски:")
            for r in h.risks:
                doc.add_paragraph(r, style="List Bullet")
        if h.test_plan:
            doc.add_paragraph("Дорожная карта проверки:")
            for step in h.test_plan:
                doc.add_paragraph(step, style="List Number")
        if h.required_resources:
            doc.add_paragraph(f"Требуемые ресурсы: {h.required_resources}")
        if h.citations:
            doc.add_paragraph("Источники:")
            for c in h.citations:
                doc.add_paragraph(f"[{c['n']}] {c['source']}: «{c['quote'][:200]}…»",
                                  style="List Bullet")

    doc.add_heading("3. Методика", level=1)
    doc.add_paragraph(
        "Гипотезы сгенерированы системой «Фабрика гипотез»: диагностические "
        "правила квантифицируют потери по данным отчёта института (классы "
        "крупности × минералогия), граф знаний связывает формы потерь с "
        "управляемыми параметрами и механизмами, языковая модель (YandexGPT) "
        "формулирует проверяемые гипотезы строго на основе извлечённого "
        "контекста и литературы. Ранжирование — взвешенная сумма пяти "
        "объяснимых критериев; вклад каждого критерия приведён в разборе "
        "оценки.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ------------------------------------------------------------------ HTML
def export_html(result: PipelineResult) -> str:
    """Автономный HTML-отчёт: открывается в браузере без сервера."""
    e = html_mod.escape

    def fmt(x):
        return "—" if x is None else f"{x:,.0f}".replace(",", " ")

    parts = [f"""<!DOCTYPE html><html lang="ru"><head><meta charset="utf-8">
<title>Фабрика гипотез — {e(result.report.plant)}</title><style>
body{{font:15px/1.55 "Segoe UI",sans-serif;color:#1c2330;max-width:960px;margin:24px auto;padding:0 16px}}
h1{{font-size:24px}} h2{{font-size:19px;margin-top:28px}} h3{{font-size:16px;margin-bottom:4px}}
.card{{border:1px solid #ccd4e0;border-radius:10px;padding:14px 18px;margin:12px 0}}
.muted{{color:#5b6878;font-size:13px}}
.sig{{border-left:4px solid #e8a33d;padding:6px 12px;margin:8px 0;background:#faf6ee}}
table{{border-collapse:collapse;font-size:13px;margin:6px 0}}
td,th{{border:1px solid #ccd4e0;padding:4px 10px;text-align:left}}
.score{{font-weight:700;background:#1f6feb;color:#fff;border-radius:12px;padding:1px 10px;float:right}}
ol,ul{{margin:4px 0 4px 22px}}
.cite{{background:#f4f6f9;border:1px solid #dde3ec;border-radius:6px;padding:6px 10px;font-size:12px;margin:5px 0}}
</style></head><body>
<h1>Фабрика гипотез — отчёт</h1>
<p><b>Цель:</b> {e(result.goal)}<br><b>Ограничения:</b> {e(result.constraints)}<br>
<b>Источник данных:</b> {e(result.report.source_file)} ·
<b>Время генерации:</b> {result.elapsed_sec:.0f} с ·
<b>База знаний:</b> {result.kb_stats.get('literature_chunks', 0)} фрагментов литературы,
граф {result.kb_stats.get('nodes', 0)} узлов / {result.kb_stats.get('edges', 0)} связей</p>

<h2>1. Диагностика потерь (факты из отчёта по хвостам)</h2>
<p class="muted">Каждый сигнал — прозрачное правило над данными отчёта:
где и в какой минеральной форме теряется металл, в тоннах за период.</p>"""]
    for s in result.signals:
        parts.append(f'<div class="sig"><b>{e(s.title)}</b> — {fmt(s.tonnes)} т '
                     f'{e(s.metal)} ({s.share_of_stream_pct:.0f}% потерь потока '
                     f'«хвосты {e(s.stream)}»)<br>'
                     f'<span class="muted">{e(s.explanation)}</span></div>')

    parts.append("<h2>2. Ранжированные гипотезы</h2>")
    for h in result.hypotheses:
        eff = h.expected_effect
        rows = "".join(
            f"<tr><td>{CRITERIA_RU.get(k, k)}</td><td>{sc['score']:.2f}</td>"
            f"<td>{(sc.get('weight') or 0):.2f}</td><td>{e(sc['explanation'])}</td></tr>"
            for k, sc in h.scores.items())
        risks = "".join(f"<li>{e(r)}</li>" for r in h.risks)
        plan = "".join(f"<li>{e(t)}</li>" for t in h.test_plan)
        cites = "".join(f'<div class="cite">[{c["n"]}] {e(c["source"])} — '
                        f'«{e(c["quote"][:250])}…»</div>' for c in h.citations)
        parts.append(f"""<div class="card">
<span class="score">{h.total_score:.2f}</span>
<h3>№{h.rank}. {e(h.statement)}</h3>
<div class="muted">поток: хвосты {e(h.stream)} · металл: {e(h.metal)} ·
сигнал: {e(h.signal_title)} ({fmt(h.signal_tonnes)} т)</div>
<p><b>Механизм:</b> {e(h.mechanism)}<br><b>Обоснование:</b> {e(h.rationale)}<br>
<b>Ожидаемый эффект:</b> возврат {fmt(eff['tonnes_min'])}–{fmt(eff['tonnes_max'])} т
{e(h.metal)}/период ({eff['min_recovery_pct']:.0f}–{eff['max_recovery_pct']:.0f}%
{e(eff['basis'])})</p>
<table><tr><th>Критерий</th><th>Балл</th><th>Вес</th><th>Почему</th></tr>{rows}</table>
<b>Риски:</b><ul>{risks}</ul>
<b>Дорожная карта проверки:</b><ol>{plan}</ol>
<b>Ресурсы:</b> {e(h.required_resources or "—")}
{cites}</div>""")

    parts.append("</body></html>")
    return "\n".join(parts)


# ------------------------------------------------------------------ CSV
def export_csv(result: PipelineResult) -> str:
    buf = io.StringIO()
    w = csv.writer(buf, delimiter=";")
    w.writerow(["rank", "score", "statement", "mechanism", "stream", "metal",
                "signal", "signal_tonnes", "effect_tonnes_min",
                "effect_tonnes_max", "lever", "capex_class", "risks",
                "test_plan", "sources"])
    for h in result.hypotheses:
        w.writerow([
            h.rank, f"{h.total_score:.3f}", h.statement, h.mechanism,
            h.stream, h.metal, h.signal_title, f"{h.signal_tonnes:.0f}",
            f"{h.expected_effect['tonnes_min']:.0f}",
            f"{h.expected_effect['tonnes_max']:.0f}",
            h.lever_label, h.capex_class,
            " | ".join(h.risks), " | ".join(h.test_plan),
            " | ".join(c["source"] for c in h.citations)])
    return buf.getvalue()


# ------------------------------------------------------------------ JSON (задачи)
def export_tasks_json(result: PipelineResult) -> str:
    """Формат, пригодный для импорта в Jira/YouTrack (bulk create)."""
    tasks = []
    for h in result.hypotheses:
        tasks.append({
            "summary": f"[Гипотеза №{h.rank}] {h.statement[:200]}",
            "description": (
                f"{h.rationale}\n\nМеханизм: {h.mechanism}\n"
                f"Ожидаемый эффект: {h.expected_effect['tonnes_min']:,.0f}–"
                f"{h.expected_effect['tonnes_max']:,.0f} т {h.metal}\n\n"
                "План проверки:\n" + "\n".join(f"- {s}" for s in h.test_plan)
                + "\n\nРиски:\n" + "\n".join(f"- {r}" for r in h.risks)),
            "labels": ["hypothesis-factory", h.stream, h.metal],
            "priority": "High" if h.rank <= 3 else "Medium",
            "customFields": {
                "score": h.total_score,
                "signal": h.signal_title,
                "sources": [c["source"] for c in h.citations],
            },
        })
    return json.dumps({"issues": tasks}, ensure_ascii=False, indent=1)
